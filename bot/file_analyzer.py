"""
File analysis module.
Handles extraction of text content from various file types:
PDF (PyMuPDF/fitz), DOCX (python-docx), XLSX (openpyxl),
code files (plain text), ZIP archives (zipfile).
"""

import asyncio
import io
import logging
import os
import zipfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Maximum text length to extract (to avoid overwhelming the AI)
MAX_EXTRACT_LENGTH = 15000

# Code file extensions that can be read as plain text
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css", ".scss",
    ".java", ".cpp", ".c", ".h", ".hpp", ".cs", ".go", ".rs",
    ".rb", ".php", ".swift", ".kt", ".scala", ".lua", ".r",
    ".sql", ".sh", ".bash", ".zsh", ".bat", ".ps1",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".xml", ".svg", ".md", ".txt", ".log", ".csv", ".env",
    ".dockerfile", ".makefile", ".gitignore", ".editorconfig",
}

# Text-based extensions (broader set)
TEXT_EXTENSIONS = CODE_EXTENSIONS | {
    ".rst", ".tex", ".bib", ".org", ".adoc",
}


def _truncate_text(text: str, max_length: int = MAX_EXTRACT_LENGTH) -> str:
    """Truncate text to max_length with an indicator."""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "\n\n... [matn qisqartirildi, jami belgilar soni: {}]".format(len(text))


async def analyze_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file using PyMuPDF (fitz).

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content or error message
    """
    try:
        import fitz  # PyMuPDF

        def _read_pdf():
            doc = fitz.open(file_path)
            text_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"--- Sahifa {page_num + 1} ---\n{page_text}")

            doc.close()
            return text_parts

        loop = asyncio.get_event_loop()
        text_parts = await loop.run_in_executor(None, _read_pdf)

        if not text_parts:
            return "PDF fayldan matn ajratib olib bo'lmadi (rasm asosidagi PDF bo'lishi mumkin)."

        full_text = "\n\n".join(text_parts)
        return _truncate_text(full_text)

    except ImportError:
        return "PDF tahlil qilish uchun PyMuPDF kutubxonasi o'rnatilmagan."
    except Exception as e:
        logger.error(f"PDF analysis error: {e}")
        return f"PDF faylni tahlil qilishda xatolik: {str(e)}"


async def analyze_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content or error message
    """
    try:
        from docx import Document

        def _read_docx():
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return text_parts

        loop = asyncio.get_event_loop()
        text_parts = await loop.run_in_executor(None, _read_docx)

        if not text_parts:
            return "DOCX fayldan matn ajratib olib bo'lmadi."

        full_text = "\n".join(text_parts)
        return _truncate_text(full_text)

    except ImportError:
        return "DOCX tahlil qilish uchun python-docx kutubxonasi o'rnatilmagan."
    except Exception as e:
        logger.error(f"DOCX analysis error: {e}")
        return f"DOCX faylni tahlil qilishda xatolik: {str(e)}"


async def analyze_xlsx(file_path: str) -> str:
    """
    Extract text content from an XLSX file using openpyxl.

    Args:
        file_path: Path to the XLSX file

    Returns:
        Extracted text content or error message
    """
    try:
        from openpyxl import load_workbook

        def _read_xlsx():
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"--- Varaq: {sheet_name} ---")

                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(v for v in row_values if v)
                    if row_text:
                        text_parts.append(row_text)
                        row_count += 1
                        if row_count > 200:  # Limit rows per sheet
                            text_parts.append("... [qolgan qatorlar qisqartirildi]")
                            break

            wb.close()
            return text_parts

        loop = asyncio.get_event_loop()
        text_parts = await loop.run_in_executor(None, _read_xlsx)

        if not text_parts:
            return "XLSX fayldan ma'lumot ajratib olib bo'lmadi."

        full_text = "\n".join(text_parts)
        return _truncate_text(full_text)

    except ImportError:
        return "XLSX tahlil qilish uchun openpyxl kutubxonasi o'rnatilmagan."
    except Exception as e:
        logger.error(f"XLSX analysis error: {e}")
        return f"XLSX faylni tahlil qilishda xatolik: {str(e)}"


async def analyze_code_file(file_path: str) -> str:
    """
    Read a code/text file as plain text.

    Args:
        file_path: Path to the code file

    Returns:
        File content or error message
    """
    try:
        def _read_code():
            # Try UTF-8 first, then fallback to latin-1
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
            return content

        loop = asyncio.get_event_loop()
        content = await loop.run_in_executor(None, _read_code)

        if not content.strip():
            return "Fayl bo'sh."

        ext = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        header = f"--- Fayl: {file_name} (turi: {ext}) ---\n"

        return header + _truncate_text(content)

    except Exception as e:
        logger.error(f"Code file analysis error: {e}")
        return f"Faylni o'qishda xatolik: {str(e)}"


async def analyze_zip(file_path: str) -> str:
    """
    List contents of a ZIP archive and extract text from readable files inside.

    Args:
        file_path: Path to the ZIP file

    Returns:
        Archive contents description or error message
    """
    try:
        def _read_zip():
            text_parts = []
            text_parts.append("--- ZIP arxiv tarkibi ---")

            with zipfile.ZipFile(file_path, "r") as zf:
                file_list = zf.namelist()
                text_parts.append(f"Jami fayllar soni: {len(file_list)}\n")

                # List all files
                text_parts.append("Fayllar ro'yxati:")
                for name in file_list[:50]:  # Show max 50 files
                    info = zf.getinfo(name)
                    size_kb = info.file_size / 1024
                    text_parts.append(f"  - {name} ({size_kb:.1f} KB)")

                if len(file_list) > 50:
                    text_parts.append(f"  ... va yana {len(file_list) - 50} ta fayl")

                # Try to read small text files from the archive
                text_parts.append("\n--- Matn fayllarining mazmuni ---")
                extracted_count = 0
                for name in file_list:
                    if extracted_count >= 5:  # Limit to 5 files
                        break
                    ext = Path(name).suffix.lower()
                    if ext in TEXT_EXTENSIONS and not name.endswith("/"):
                        info = zf.getinfo(name)
                        if info.file_size < 50000:  # Only files < 50KB
                            try:
                                content = zf.read(name).decode("utf-8", errors="replace")
                                if content.strip():
                                    text_parts.append(f"\n--- {name} ---")
                                    text_parts.append(content[:3000])
                                    extracted_count += 1
                            except Exception:
                                pass

            return text_parts

        loop = asyncio.get_event_loop()
        text_parts = await loop.run_in_executor(None, _read_zip)

        full_text = "\n".join(text_parts)
        return _truncate_text(full_text)

    except zipfile.BadZipFile:
        return "Bu fayl yaroqli ZIP arxiv emas."
    except Exception as e:
        logger.error(f"ZIP analysis error: {e}")
        return f"ZIP arxivni tahlil qilishda xatolik: {str(e)}"


async def analyze_file(file_path: str, file_name: str) -> str:
    """
    Analyze a file based on its extension.
    Routes to the appropriate analyzer.

    Args:
        file_path: Path to the downloaded file
        file_name: Original file name (for extension detection)

    Returns:
        Extracted/analyzed text content
    """
    ext = Path(file_name).suffix.lower()

    if ext == ".pdf":
        return await analyze_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return await analyze_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        return await analyze_xlsx(file_path)
    elif ext == ".zip":
        return await analyze_zip(file_path)
    elif ext in TEXT_EXTENSIONS:
        return await analyze_code_file(file_path)
    else:
        # Try to read as text anyway
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            if content.strip():
                return _truncate_text(f"--- Fayl: {file_name} ---\n{content}")
        except (UnicodeDecodeError, Exception):
            pass

        return (
            f"'{file_name}' fayl turini tahlil qilish imkoni yo'q. "
            f"Qo'llab-quvvatlanadigan formatlar: PDF, DOCX, XLSX, ZIP, "
            f"va barcha matn/kod fayllar (py, js, html, css, xml, cpp va boshqalar)."
        )
